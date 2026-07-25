import io
from docx import Document
from docx.shared import RGBColor, Pt, Inches

def generuj_docx(name, position, description, exp_list, edu_list, skills_list, langs_list, cert_list=None, phone=None, email=None, location=None, linkedin=None, github=None, rodo=True, photo_bytes=None):
    doc = Document()
    
    if photo_bytes:
        doc.add_picture(io.BytesIO(photo_bytes), width=Inches(1.2))
        
    if name:
        doc.add_heading(name, level=0)
    if position:
        p_pos = doc.add_paragraph()
        p_pos.add_run(position).italic = True
        
    bloki_kontaktowe = []
    if phone: bloki_kontaktowe.append(f"Tel: {phone}")
    if email: bloki_kontaktowe.append(f"Email: {email}")
    if location: bloki_kontaktowe.append(f"Lokalizacja: {location}")
    if linkedin: bloki_kontaktowe.append(f"LinkedIn: {linkedin}")
    if github: bloki_kontaktowe.append(f"GitHub: {github}")
    
    if bloki_kontaktowe:
        p_kon = doc.add_paragraph()
        run_kon = p_kon.add_run(" | ".join(bloki_kontaktowe))
        run_kon.font.color.rgb = RGBColor(100, 100, 100)

    if description:
        doc.add_paragraph(description)

    if exp_list:
        doc.add_heading("Doświadczenie zawodowe", level=1)
        for job in exp_list:
            naglowek_pracy = f"{job['role']} w {job['company']}"
            if job['years']: naglowek_pracy += f" ({job['years']})"
            doc.add_heading(naglowek_pracy, level=2)
            if job['duty']:
                doc.add_paragraph(job['duty'])

    if edu_list:
        doc.add_heading("Edukacja", level=1)
        for edu in edu_list:
            naglowek_szkoły = edu['school']
            if edu['years_edu']: naglowek_szkoły += f" ({edu['years_edu']})"
            doc.add_heading(naglowek_szkoły, level=2)
            if edu['field']:
                doc.add_paragraph(f"Kierunek: {edu['field']}")

    if skills_list:
        doc.add_heading("Umiejętności", level=1)
        tekst_skilli = []

        for s in skills_list:
            if s["level"] is None:
                tekst_skilli.append(s["skill"])
            else:
                gwiazdki = "★" * s["level"] + "☆" * (5 - s["level"])
                tekst_skilli.append(f"{s['skill']} ({gwiazdki})")

        doc.add_paragraph(", ".join(tekst_skilli))

    if langs_list:
        doc.add_heading("Języki obce", level=1)
        for lang in langs_list:
            doc.add_paragraph(f"{lang['lang']} - {lang['level']}")

    if cert_list:
        doc.add_heading("Certyfikaty i Kursy", level=1)
        for cert in cert_list:
            doc.add_paragraph(f"- {cert}")

    if rodo:
        doc.add_paragraph()  
        p_rodo = doc.add_paragraph()
        p_rodo.alignment = 1
        run_rodo = p_rodo.add_run("Wyrażam zgodę na przetwarzanie moich danych osobowych dla potrzeb niezbędnych do realizacji procesu rekrutacji (zgodnie z rozporządzeniem o ochronie danych osobowych RODO).")
        run_rodo.font.size = Pt(8.5)
        run_rodo.font.color.rgb = RGBColor(120, 120, 120)
            
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()